import scrapy
import urllib.parse
from pathlib import Path
from docx import Document

class NorthAmericanAnimalsSpider(scrapy.Spider):
	name = 'north_american_animals'

	def start_requests(self):
		urls = [
			'https://www.wildlifenorthamerica.com/A-Z/Mammal/common.html',
			'https://www.wildlifenorthamerica.com/A-Z/Reptile/common.html',
		]
		for url in urls:
			yield scrapy.Request(url=url, callback=self.parse)

	def parse(self, response):
		length_animal_list = len(response.xpath('/html/body/div/table[2]/td/table/tr').getall())
		for i in range(2, length_animal_list + 1):
			animal_name = response.xpath(f'/html/body/div/table[2]/td/table/tr[{i}]/td[1]/a/text()').get()
			encoded_animal_name = urllib.parse.quote(animal_name)
			wiki_url = f'https://en.wikipedia.org/wiki/{encoded_animal_name}'
			yield scrapy.Request(url=wiki_url, callback=self.parse_wiki_page, meta={'animal': animal_name})

	def parse_wiki_page(self, response):
		animal_description = response.xpath('string(//*[@id="mw-content-text"]/div/p[2])').get()
		animal_name = response.meta.get('animal')
		Path("descriptions").mkdir(parents=True, exist_ok=True)
		filename = f'descriptions/{animal_name}.docx'

		document = Document()
		document.add_heading(animal_name, 0)
		document.add_paragraph(animal_description)
		document.save(filename)
		self.log(f'Saved file {filename}')
